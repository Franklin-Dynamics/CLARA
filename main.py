import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import csv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import re
import json

class CompanyFilingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Franklin Dynamics - C.L.A.R.A. (Company Legal & Accounting Reporting Assistant)")

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(padx=10, pady=10, expand=True, fill="both")

        self.company_info_vars = {}
        self.cs01_vars = {}
        self.accounts_vars = {}
        self.ct600_vars = {}
        self.expenses_entries = []
        self.logo_path = None

        self.create_company_info_tab()
        self.create_cs01_tab()
        self.create_expenses_tab()
        self.create_accounts_tab()
        self.create_ct600_tab()

        btn_frame = ttk.Frame(root)
        btn_frame.pack(pady=10)

        ttk.Button(btn_frame, text="Submit & Download DOCX", command=self.submit).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="Import Metrics CSV", command=self.import_metrics_csv).grid(row=0, column=1, padx=5)
        ttk.Button(btn_frame, text="Load Previous Data", command=self.load_data).grid(row=0, column=2, padx=5)
        ttk.Button(btn_frame, text="Choose Logo", command=self.choose_logo).grid(row=0, column=3, padx=5)

    def create_form(self, parent, fields, vars_dict):
        for idx, (label, key) in enumerate(fields):
            ttk.Label(parent, text=label).grid(row=idx, column=0, sticky="w", padx=5, pady=5)
            entry = ttk.Entry(parent)
            entry.grid(row=idx, column=1, sticky="ew", padx=5, pady=5)
            vars_dict[key] = entry
        parent.columnconfigure(1, weight=1)

    def create_company_info_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Company Info")

        fields = [
            ("Company Name", "company_name"),
            ("Company Number", "company_number"),
            ("Email", "email"),
        ]
        self.create_form(tab, fields, self.company_info_vars)

    def create_cs01_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="CS01 Confirmation Statement")

        fields = [
            ("Confirmation Date (DD-MM-YYYY)", "confirmation_date"),
            ("SIC Code", "sic_code"),
            ("Shareholder Names (comma separated)", "shareholders"),
            ("Lawful Statement (Yes/No)", "lawful_statement"),
            ("Company's CORA Code", "unique_code"),
        ]
        self.create_form(tab, fields, self.cs01_vars)

    def create_expenses_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Expenses")

        self.expenses_container = ttk.Frame(tab)
        self.expenses_container.pack(fill="both", expand=True, padx=5, pady=5)

        self.add_expense_row()

        btn_frame = ttk.Frame(tab)
        btn_frame.pack(pady=5)
        ttk.Button(btn_frame, text="Add Expense", command=self.add_expense_row).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="Calculate Total Expenses", command=self.calculate_expenses).pack(side="left", padx=5)

        self.total_expenses_var = tk.StringVar(value="0.00")
        ttk.Label(tab, text="Total Expenses:").pack(side="left", padx=5)
        ttk.Label(tab, textvariable=self.total_expenses_var).pack(side="left")

    def add_expense_row(self):
        frame = ttk.Frame(self.expenses_container)
        frame.pack(fill="x", pady=2)

        description_entry = ttk.Entry(frame, width=40)
        description_entry.pack(side="left", padx=5)
        amount_entry = ttk.Entry(frame, width=15)
        amount_entry.pack(side="left", padx=5)

        image_path_var = tk.StringVar(value="")  # Store path to uploaded image for this expense

        def upload_image():
            path = filedialog.askopenfilename()
            if path:
                image_path_var.set(path)
                messagebox.showinfo("Image Selected", f"Receipt image set:\n{os.path.basename(path)}")

        upload_btn = ttk.Button(frame, text="Upload Image", command=upload_image)
        upload_btn.pack(side="left", padx=5)

        # Store all 3 widgets/data: description, amount, image_path_var
        self.expenses_entries.append((description_entry, amount_entry, image_path_var))

    def calculate_expenses(self):
        total = 0.0
        for desc_entry, amt_entry in self.expenses_entries:
            amt_str = amt_entry.get().strip()
            if amt_str:
                try:
                    amt = float(amt_str)
                    total += amt
                except ValueError:
                    messagebox.showerror("Invalid Input", f"Invalid expense amount: {amt_str}")
                    return
        self.total_expenses_var.set(f"{total:.2f}")
        self.update_profit_loss(total)

    def update_profit_loss(self, total_expenses):
        turnover_entry = self.accounts_vars.get("turnover")
        if turnover_entry:
            try:
                turnover = float(turnover_entry.get().strip())
                profit_loss = turnover - total_expenses
                if "profit_loss" in self.accounts_vars:
                    self.accounts_vars["profit_loss"].delete(0, tk.END)
                    self.accounts_vars["profit_loss"].insert(0, f"{profit_loss:.2f}")
            except ValueError:
                messagebox.showerror("Invalid Input", "Turnover must be a valid number.")

    def create_accounts_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Annual Accounts")

        fields = [
            ("Account Date (DD-MM-YYYY)", "account_date"),
            ("Turnover", "turnover"),
            ("Profit/Loss (auto-calculated)", "profit_loss"),
            ("Assets Value", "assets"),
            ("Liabilities Value", "liabilities"),
        ]
        self.create_form(tab, fields, self.accounts_vars)

        if "turnover" in self.accounts_vars:
            self.accounts_vars["turnover"].bind("<FocusOut>", lambda e: self.calculate_expenses())

    def create_ct600_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Corporation Tax Return (CT600)")

        fields = [
            ("Tax Year Start (YYYY-MM-DD)", "tax_year_start"),
            ("Tax Year End (YYYY-MM-DD)", "tax_year_end"),
            ("Profit Before Tax", "profit_before_tax"),
            ("Tax Payable", "tax_payable"),
        ]
        self.create_form(tab, fields, self.ct600_vars)

    def gather_data(self):
        shared_info = {k: v.get().strip() for k, v in self.company_info_vars.items()}
        cs01 = {k: v.get().strip() for k, v in self.cs01_vars.items()}
        accounts = {k: v.get().strip() for k, v in self.accounts_vars.items()}
        ct600 = {k: v.get().strip() for k, v in self.ct600_vars.items()}

        shareholders = cs01.get("shareholders", "")
        cs01["shareholders"] = [s.strip() for s in shareholders.split(",") if s.strip()]

        expenses_list = [
            {
                "description": desc.get(),
                "amount": amt.get(),
                "receipt_image": img_path.get()  # add this line
            }
            for desc, amt, img_path in self.expenses_entries if amt.get().strip()
        ]

        return {
            "company_info": shared_info,
            "cs01": cs01,
            "accounts": accounts,
            "ct600": ct600,
            "expenses": expenses_list,
            "total_expenses": self.total_expenses_var.get(),
            "logo": self.logo_path
        }

    def choose_logo(self):
        path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg")])
        if path:
            self.logo_path = path

    def submit(self):
        data = self.gather_data()
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if not save_path:
            return

        try:
            doc = Document()

            if data["logo"]:
                doc.add_picture(data["logo"], width=Inches(2))

            doc.add_heading('Annual Company Filing Report', 0)

            doc.add_heading('Company Information', level=1)
            for key, value in data['company_info'].items():
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

            doc.add_heading('Confirmation Statement (CS01)', level=1)
            cs01 = data['cs01']
            doc.add_paragraph(f"Confirmation Date: {cs01['confirmation_date']}")
            doc.add_paragraph(f"SIC Code: {cs01['sic_code']}")
            doc.add_paragraph(f"Lawful Statement: {cs01['lawful_statement']}")
            doc.add_paragraph(f"CORA Code: {cs01['unique_code']}")
            doc.add_paragraph("Shareholders:")
            for shareholder in cs01['shareholders']:
                doc.add_paragraph(f"• {shareholder}", style='List Bullet')

            doc.add_heading('Expenses', level=1)
            for expense in data['expenses']:
                doc.add_paragraph(f"{expense['description']}: £{expense['amount']}")
                if expense.get('receipt_image'):
                    try:
                        doc.add_picture(expense['receipt_image'], width=Inches(2))
                    except Exception as e:
                        doc.add_paragraph(f"[Failed to load receipt image: {expense['receipt_image']}]")
            doc.add_paragraph(f"Total Expenses: £{data['total_expenses']}")

            doc.add_heading('Annual Accounts', level=1)
            acc = data['accounts']
            doc.add_paragraph(f"Account Date: {acc['account_date']}")
            doc.add_paragraph(f"Turnover: £{acc['turnover']}")
            doc.add_paragraph(f"Profit/Loss: £{acc['profit_loss']}")
            doc.add_paragraph(f"Assets Value: £{acc['assets']}")
            doc.add_paragraph(f"Liabilities Value: £{acc['liabilities']}")

            doc.add_heading('Corporation Tax Return (CT600)', level=1)
            ct600 = data['ct600']
            doc.add_paragraph(f"Tax Year Start: {ct600['tax_year_start']}")
            doc.add_paragraph(f"Tax Year End: {ct600['tax_year_end']}")
            doc.add_paragraph(f"Profit Before Tax: £{ct600['profit_before_tax']}")
            doc.add_paragraph(f"Tax Payable: £{ct600['tax_payable']}")

            doc.save(save_path)

            with open("last_data.json", "w") as f:
                json.dump(data, f)

            messagebox.showinfo("Success", f"Document saved to {save_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate DOCX: {e}")

    def import_metrics_csv(self):
        path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
        if not path:
            return
        try:
            with open(path, newline='') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            if rows:
                last = rows[-1]
                for k in self.accounts_vars:
                    if k in last:
                        self.accounts_vars[k].delete(0, tk.END)
                        self.accounts_vars[k].insert(0, last[k])
                messagebox.showinfo("Success", "Imported metrics from CSV.")
            else:
                messagebox.showwarning("Warning", "CSV is empty.")
        except Exception as e:
            messagebox.showerror("Error", f"Import failed: {e}")

    def load_data(self):
        try:
            with open("last_data.json") as f:
                data = json.load(f)

            for k, v in data.get("company_info", {}).items():
                if k in self.company_info_vars:
                    self.company_info_vars[k].delete(0, tk.END)
                    self.company_info_vars[k].insert(0, v)

            for k, v in data.get("cs01", {}).items():
                if isinstance(v, list):
                    v = ", ".join(v)
                if k in self.cs01_vars:
                    self.cs01_vars[k].delete(0, tk.END)
                    self.cs01_vars[k].insert(0, v)

            for k, v in data.get("accounts", {}).items():
                if k in self.accounts_vars:
                    self.accounts_vars[k].delete(0, tk.END)
                    self.accounts_vars[k].insert(0, v)

            for k, v in data.get("ct600", {}).items():
                if k in self.ct600_vars:
                    self.ct600_vars[k].delete(0, tk.END)
                    self.ct600_vars[k].insert(0, v)

            self.logo_path = data.get("logo")
            messagebox.showinfo("Success", "Previous data loaded.")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load data: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = CompanyFilingApp(root)
    root.mainloop()